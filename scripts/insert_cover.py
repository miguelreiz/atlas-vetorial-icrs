import sys
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_image_to_docx(doc_path, img_path, out_path):
    doc = docx.Document(doc_path)
    
    # We want to insert the image right at the beginning of the document.
    # The first paragraph usually has the padding/title.
    # Let's insert a new paragraph before the first paragraph.
    p = doc.paragraphs[0].insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run()
    # Add the image. We'll make it 5 inches wide to fit nicely on the page
    run.add_picture(img_path, width=Inches(6.0))
    
    # Add a bit of space below the image
    p.paragraph_format.space_after = Pt(24)
    
    doc.save(out_path)
    print(f"Cover added successfully: {out_path}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: script.py <input_docx> <image_path> <output_docx>")
        sys.exit(1)
        
    insert_image_to_docx(sys.argv[1], sys.argv[2], sys.argv[3])
