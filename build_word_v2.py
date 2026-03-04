#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_word_v2.py — Atlas Vetorial ICRS
Word profissional: imagens no lugar certo, tabelas, headings hierárquicos.
"""
import os, re, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = Path(__file__).parent
CHDIR = BASE / "chapters" / "pt"
OUT   = BASE / "Atlas_Vetorial_ICRS_v2.docx"

CHAPTERS = [
    "CH-000_Prefacio.md",
    "CH-001_Anatomia_Corneana.md",
    "CH-002_Biomecanica_Aneis.md",
    "CH-003_Classificacao_Ceratocone.md",
    "CH-004_VR_Vetor_Radial.md",
    "CH-005_VT_Vetor_Tangencial.md",
    "CH-006_Vt_Vetor_Torque.md",
    "CH-007_VComa_Deslocamento_Optico.md",
    "CH-008_LDM_A_Lei_do_Disco_Mecanico.md",
    "CH-009_VEsferico_Soma_Vetorial.md",
    "CH-010_ICE_Indice_Coerencia_Eixos.md",
    "CH-011_Nomogramas_Vetoriais.md",
    "CH-012_Casos_Clinicos.md",
    "CH-013_Complicacoes_Manejo.md",
    "CH-014_Futuro_Aneis.md",
    "ENCARTE_Referencia_Rapida.md",
    "REFERENCIAS_BIBLIOGRAFICAS.md",
]

# ── Cores ──────────────────────────────────────────────────────────────────
C_H1   = RGBColor(0x0D, 0x2B, 0x4D)
C_H2   = RGBColor(0x1A, 0x5C, 0x8A)
C_H3   = RGBColor(0x2E, 0x86, 0xC1)
C_H4   = RGBColor(0x5D, 0x6D, 0x7E)
C_BQ   = RGBColor(0x2E, 0x4F, 0x6E)
C_CAP  = RGBColor(0x5D, 0x6D, 0x7E)
C_HEAD = "D6E4F0"    # azul claro p/ cabeçalho de tabela
C_ALT  = "EBF5FB"    # azul muito claro p/ linhas alternadas

# ── Utilitários ──────────────────────────────────────────────────────────
def strip_md_inline(t):
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'\*(.+?)\*',     r'\1', t)
    t = re.sub(r'`(.+?)`',       r'\1', t)
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    return t.strip()

def shd_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def find_image(raw):
    raw = raw.strip().strip('"')
    if os.path.isabs(raw):
        p = Path(raw);  return p if p.exists() else None
    p = BASE / raw;     return p if p.exists() else None

def run_with_fmt(para, text):
    """Adiciona runs com bold/italic/code a um parágrafo existente."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'; r.font.size = Pt(9.5)
        else:
            para.add_run(part)

# ── Estilos ──────────────────────────────────────────────────────────────
def setup_styles(doc):
    def set_h(name, size, color, bold=True):
        s = doc.styles[name]
        s.font.name  = 'Calibri Light'
        s.font.size  = Pt(size)
        s.font.bold  = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)

    set_h('Heading 1', 24, C_H1)
    set_h('Heading 2', 17, C_H2)
    set_h('Heading 3', 13, C_H3)
    set_h('Heading 4', 11, C_H4, bold=False)

    n = doc.styles['Normal']
    n.font.name = 'Calibri'; n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = Pt(14)

# ── Capa ─────────────────────────────────────────────────────────────────
def add_cover(doc):
    for _ in range(6): doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ATLAS VETORIAL ICRS")
    r.font.name = 'Calibri Light'; r.font.size = Pt(38)
    r.bold = True; r.font.color.rgb = C_H1

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("O Sistema Vetorial de Planejamento de\nAneis Intracorneanos")
    r2.font.name = 'Calibri'; r2.font.size = Pt(18)
    r2.font.color.rgb = C_H2

    doc.add_paragraph()
    e = doc.add_paragraph()
    e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = e.add_run("Edicao 2026 — Versao Draft")
    r3.font.name = 'Calibri'; r3.font.size = Pt(13)
    r3.italic = True; r3.font.color.rgb = C_H4

    doc.add_paragraph()
    la = doc.add_paragraph()
    la.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = la.add_run("Portugues Brasileiro")
    r4.font.name = 'Calibri'; r4.font.size = Pt(11)

    doc.add_page_break()

# ── Tabela ───────────────────────────────────────────────────────────────
def parse_md_table(lines):
    rows = []
    for ln in lines:
        if re.match(r'\s*\|[-:| ]+\|\s*$', ln): continue
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

def add_table(doc, rows):
    if not rows: return
    nc = max(len(r) for r in rows)
    rows = [r + ['']*(nc-len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=nc)
    tbl.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = tbl.cell(i, j)
            clean = strip_md_inline(txt)
            p = cell.paragraphs[0]
            rn = p.add_run(clean)
            if i == 0:
                rn.bold = True
                shd_cell(cell, C_HEAD)
            elif i % 2 == 0:
                shd_cell(cell, C_ALT)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

# ── Parser principal ──────────────────────────────────────────────────────
def process(doc, path: Path, stats: dict):
    raw = path.read_text(encoding='utf-8')

    # Remove blocos YAML e metadados
    raw = re.sub(r'^```yaml.*?```\s*', '', raw, flags=re.DOTALL|re.MULTILINE)
    raw = re.sub(r'^---\s*$', '', raw, flags=re.MULTILINE)
    raw = re.sub(r'^\*Pipeline Status:.*$', '', raw, flags=re.MULTILINE)

    lines = raw.splitlines()
    i = 0
    tbl_buf  = []
    in_tbl   = False
    code_buf = []
    in_code  = False

    while i < len(lines):
        ln = lines[i]

        # ── Bloco de código ───────────────────────────────────────────────
        if in_code:
            if ln.strip().startswith('```'):
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Courier New'; run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1.2)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(6)
                # fundo cinza via shading no parágrafo via xml
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear')
                shd.set(qn('w:color'),'auto')
                shd.set(qn('w:fill'),'F2F3F4')
                pPr.append(shd)
                code_buf = []
            else:
                code_buf.append(ln)
            i += 1; continue

        if re.match(r'\s*```', ln):
            in_code = True; i += 1; continue

        # ── Tabela ────────────────────────────────────────────────────────
        if re.match(r'\s*\|', ln):
            if not in_tbl:
                in_tbl = True; tbl_buf = []
            tbl_buf.append(ln)
            i += 1; continue
        elif in_tbl:
            in_tbl = False
            add_table(doc, parse_md_table(tbl_buf))
            tbl_buf = []

        # ── Imagem ────────────────────────────────────────────────────────
        m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', ln.strip())
        if m:
            alt, raw_p = m.group(1), m.group(2)
            img = find_image(raw_p)
            if img:
                try:
                    doc.add_picture(str(img), width=Inches(5.8))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(alt[:220] if alt else '')
                    cr.italic = True; cr.font.size = Pt(9)
                    cr.font.color.rgb = C_CAP
                    cap.paragraph_format.space_after = Pt(10)
                    stats['imgs'] += 1
                except Exception as ex:
                    doc.add_paragraph(f"[Imagem: {alt}]")
            else:
                p = doc.add_paragraph(f"[Imagem nao localizada: {Path(raw_p).name}]")
                for r in p.runs: r.italic = True
            i += 1; continue

        # ── Headings ──────────────────────────────────────────────────────
        hm = re.match(r'^(#{1,6})\s+(.+)', ln)
        if hm:
            lvl  = min(len(hm.group(1)), 4)
            txt  = strip_md_inline(hm.group(2))
            doc.add_heading(txt, level=lvl)
            i += 1; continue

        # ── Blockquote ────────────────────────────────────────────────────
        bm = re.match(r'^>\s*(.*)', ln)
        if bm:
            txt = strip_md_inline(bm.group(1))
            if txt:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                r = p.add_run(txt)
                r.italic = True; r.font.color.rgb = C_BQ
            i += 1; continue

        # ── Lista ─────────────────────────────────────────────────────────
        lm = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', ln)
        if lm:
            indent = len(lm.group(1))
            txt    = lm.group(3)
            is_num = bool(re.match(r'\d+\.', lm.group(2)))
            sty    = 'List Number' if is_num else 'List Bullet'
            p = doc.add_paragraph(style=sty)
            p.paragraph_format.left_indent = Cm(indent * 0.5 + 0.5)
            run_with_fmt(p, txt)
            i += 1; continue

        # ── Linha vazia ou horizontal ──────────────────────────────────────
        if not ln.strip() or re.match(r'^[-*_]{3,}$', ln.strip()):
            i += 1; continue

        # ── Parágrafo normal ──────────────────────────────────────────────
        p = doc.add_paragraph()
        run_with_fmt(p, ln)
        i += 1

    # Flush bufferes
    if in_tbl and tbl_buf:
        add_table(doc, parse_md_table(tbl_buf))

# ── Main ──────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin  = Cm(2.5)
    setup_styles(doc)
    add_cover(doc)

    stats = {'imgs': 0, 'chs': 0}
    for fn in CHAPTERS:
        fp = CHDIR / fn
        if not fp.exists():
            print(f"  SKIP: {fn}")
            continue
        print(f"  OK  : {fn}")
        process(doc, fp, stats)
        doc.add_page_break()
        stats['chs'] += 1

    doc.save(str(OUT))
    mb = OUT.stat().st_size / 1024 / 1024
    print(f"\nAtlas Word v2 gerado!")
    print(f"  Capitulos : {stats['chs']}")
    print(f"  Imagens   : {stats['imgs']}")
    print(f"  Arquivo   : {OUT}")
    print(f"  Tamanho   : {mb:.1f} MB")

if __name__ == '__main__':
    main()
