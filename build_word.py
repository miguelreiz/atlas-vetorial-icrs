#!/usr/bin/env python3
"""
build_word.py — Atlas Vetorial ICRS
Converte todos os capítulos Markdown em um único arquivo Word (.docx)
com imagens embutidas, tabelas formatadas e hierarquia de títulos.
"""

import os
import re
import glob
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
CHAPTERS_DIR = BASE_DIR / "chapters" / "pt"
IMAGES_DIR = BASE_DIR / "images"
OUTPUT_FILE = BASE_DIR / "Atlas_Vetorial_ICRS.docx"

# Ordem dos capítulos
CHAPTER_ORDER = [
    "CH-000_Prefacio.md",
    "CH-001_Anatomia_Corneana.md",
    "CH-002_Biomecanica_Aneis.md",
    "CH-003_Classificacao_Ceratocone.md",
    "CH-004_VR_Vetor_Radial.md",
    "CH-005_VT_Vetor_Tangencial.md",
    "CH-006_Vt_Vetor_Torque.md",
    "CH-007_VComa_Deslocamento_Optico.md",
    "CH-008_LDM_A_Lei_do_Disco_Mecanico.md",
    "CH-009_VEsferico_Soma_Vetorial.md",
    "CH-010_ICE_Indice_Coerencia_Eixos.md",
    "CH-011_Nomogramas_Vetoriais.md",
    "CH-012_Casos_Clinicos.md",
    "CH-013_Complicacoes_Manejo.md",
    "CH-014_Futuro_Aneis.md",
    "ENCARTE_Referencia_Rapida.md",
    "REFERENCIAS_BIBLIOGRAFICAS.md",
]


def setup_styles(doc):
    """Configura estilos do documento"""
    # Fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading 1 (# )
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4D)  # Navy

    # Heading 2 (## )
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    # Heading 3 (### )
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

    # Heading 4 (#### )
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def add_cover_page(doc):
    """Adiciona capa do Atlas"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ATLAS VETORIAL ICRS")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4D)
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(
        "O Sistema Vetorial de Planejamento de Anéis Intracorneanos"
    )
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run2.font.name = 'Calibri'

    doc.add_paragraph()
    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = edition.add_run("Edição 2026 — Versão Draft")
    run3.font.size = Pt(13)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_paragraph()
    doc.add_paragraph()
    lang_p = doc.add_paragraph()
    lang_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = lang_p.add_run("Português Brasileiro")
    run4.font.size = Pt(11)

    doc.add_page_break()


def resolve_image_path(img_path_raw: str) -> Path | None:
    """Resolve o caminho de uma imagem relativa ao diretório base."""
    # Remove aspas se houver
    img_path_raw = img_path_raw.strip().strip('"').strip("'")

    # Path absoluto
    if os.path.isabs(img_path_raw):
        p = Path(img_path_raw)
        return p if p.exists() else None

    # Path relativo ao BASE_DIR
    p = BASE_DIR / img_path_raw
    if p.exists():
        return p

    # Fallback: procurar só pelo nome do arquivo
    filename = Path(img_path_raw).name
    for found in BASE_DIR.rglob(filename):
        return found

    return None


def parse_table(lines: list) -> list:
    """Parse linhas de tabela Markdown e retorna lista de listas"""
    rows = []
    for line in lines:
        if re.match(r'\s*\|[-:| ]+\|\s*$', line):
            continue  # linha separadora
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Adiciona tabela formatada ao documento"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad rows
    rows = [r + [''] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            # Remove markdown bold
            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
            cell_text = re.sub(r'`(.+?)`', r'\1', cell_text)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if i == 0:
                run.bold = True
                # Sombrear header
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D6E4F0')
                tcPr.append(shd)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def clean_inline(text: str) -> str:
    """Remove markdown inline (bold, italic, code, links simples)"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def add_formatted_paragraph(doc, text: str):
    """Adiciona parágrafo com suporte a bold e italic inline"""
    p = doc.add_paragraph()
    # Processar **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Limpar links [text](url) → text
            part = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part)
            p.add_run(part)
    return p


def process_chapter(doc, filepath: Path, img_count: list):
    """Processa um capítulo Markdown e adiciona ao documento Word"""
    with open(filepath, encoding='utf-8') as f:
        content = f.read()

    # Remove blocos YAML
    content = re.sub(r'^```yaml.*?```', '', content, flags=re.DOTALL | re.MULTILINE)

    lines = content.split('\n')
    i = 0
    table_buffer = []
    in_table = False
    in_code = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Detecta fim de bloco de código
        if in_code:
            if line.strip().startswith('```'):
                in_code = False
                # Adiciona bloco de código como parágrafo mono
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
                code_buffer = []
            else:
                code_buffer.append(line)
            i += 1
            continue

        # Detecta início de bloco de código
        if line.strip().startswith('```'):
            in_code = True
            i += 1
            continue

        # Tabelas
        if re.match(r'\s*\|', line):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            rows = parse_table(table_buffer)
            add_table_to_doc(doc, rows)
            table_buffer = []

        # Quebra de página entre capítulos (---) no início do arquivo
        if line.strip() == '---':
            i += 1
            continue

        # Imagens
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path_raw = img_match.group(2)
            img_path = resolve_image_path(img_path_raw)
            if img_path and img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Caption
                    cap = doc.add_paragraph(alt_text[:200] if alt_text else '')
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.runs[0] if cap.runs else cap.add_run(alt_text[:200])
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
                    img_count[0] += 1
                except Exception as e:
                    doc.add_paragraph(f"[Imagem: {alt_text}]").italic = True
            else:
                doc.add_paragraph(f"[Imagem não encontrada: {img_path_raw}]")
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if h_match:
            level = len(h_match.group(1))
            text = clean_inline(h_match.group(2))
            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            elif level == 3:
                doc.add_heading(text, level=3)
            else:
                doc.add_heading(text, level=4)
            i += 1
            continue

        # Blockquotes (> ...)
        bq_match = re.match(r'^>\s*(.*)', line)
        if bq_match:
            text = clean_inline(bq_match.group(1))
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.right_indent = Cm(1)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x2E, 0x4F, 0x6E)
            i += 1
            continue

        # Listas
        li_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if li_match:
            indent = len(li_match.group(1))
            text = clean_inline(li_match.group(3))
            is_numbered = bool(re.match(r'\d+\.', li_match.group(2)))
            style = 'List Number' if is_numbered else 'List Bullet'
            p = doc.add_paragraph(text, style=style)
            if indent > 0:
                p.paragraph_format.left_indent = Cm(indent * 0.5 + 1)
            i += 1
            continue

        # Linha vazia
        if not line.strip():
            i += 1
            continue

        # Linha horizontal
        if re.match(r'^[-*_]{3,}$', line.strip()):
            i += 1
            continue

        # Parágrafo normal
        add_formatted_paragraph(doc, line)
        i += 1

    # Flush tabela pendente
    if in_table and table_buffer:
        rows = parse_table(table_buffer)
        add_table_to_doc(doc, rows)


def main():
    print("📄 Gerando Atlas Vetorial ICRS — Word Document")
    print(f"   Diretório: {BASE_DIR}")

    doc = Document()

    # Configurar margens (2cm)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_cover_page(doc)

    img_count = [0]
    chapters_processed = 0
    chapters_failed = []

    for chapter_file in CHAPTER_ORDER:
        filepath = CHAPTERS_DIR / chapter_file
        if not filepath.exists():
            print(f"   ⚠️  Não encontrado: {chapter_file}")
            chapters_failed.append(chapter_file)
            continue

        print(f"   ✅ Processando: {chapter_file}")
        try:
            process_chapter(doc, filepath, img_count)
            doc.add_page_break()
            chapters_processed += 1
        except Exception as e:
            print(f"   ❌ Erro em {chapter_file}: {e}")
            chapters_failed.append(chapter_file)

    doc.save(str(OUTPUT_FILE))

    print(f"\n✅ Atlas Word gerado!")
    print(f"   Capítulos: {chapters_processed}")
    print(f"   Imagens: {img_count[0]}")
    print(f"   Arquivo: {OUTPUT_FILE}")
    if chapters_failed:
        print(f"   Falhas: {chapters_failed}")
    print(f"   Tamanho: {OUTPUT_FILE.stat().st_size / 1024 / 1024:.1f} MB")


if __name__ == '__main__':
    main()
