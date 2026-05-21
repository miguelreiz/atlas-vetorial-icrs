import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Custom XML helper functions to prevent duplicate tags and guarantee clean openXML output ---

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    existing_mar = tcPr.find(qn('w:tcMar'))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# --- Inline formatting parser using a non-greedy regex splitter ---

def parse_inline_markdown(text):
    # Matches ***bold-italic***, **bold**, *italic*, _italic_, or `code`
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|_.*?_|`.*?`)')
    parts = pattern.split(text)
    
    result = []
    for part in parts:
        if not part:
            continue
        
        if part.startswith('***') and part.endswith('***'):
            result.append((part[3:-3], True, True, False))
        elif part.startswith('**') and part.endswith('**'):
            result.append((part[2:-2], True, False, False))
        elif part.startswith('*') and part.endswith('*'):
            result.append((part[1:-1], False, True, False))
        elif part.startswith('_') and part.endswith('_'):
            result.append((part[1:-1], False, True, False))
        elif part.startswith('`') and part.endswith('`'):
            result.append((part[1:-1], False, False, True))
        else:
            result.append((part, False, False, False))
            
    return result

def add_runs_to_paragraph(paragraph, parsed_parts, default_font_size=11, default_font_name="Calibri", is_blockquote=False, is_code=False):
    for text_part, is_bold, is_italic, is_inline_code in parsed_parts:
        run = paragraph.add_run(text_part)
        
        run.font.name = "Consolas" if (is_inline_code or is_code) else default_font_name
        if is_inline_code:
            run.font.size = Pt(default_font_size - 1.5)
            run.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
        else:
            run.font.size = Pt(default_font_size)
            if is_blockquote:
                run.font.color.rgb = RGBColor(0x33, 0x4E, 0x68)
            else:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                
        run.bold = is_bold
        run.italic = is_italic

def add_runs_to_heading(heading_para, parsed_parts, font_size=18, font_name="Calibri"):
    for text_part, is_bold, is_italic, is_inline_code in parsed_parts:
        run = heading_para.add_run(text_part)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy blue for all headings
        run.bold = is_bold
        run.italic = is_italic

# --- Heading styles setup ---

def configure_heading_style(doc, style_name, font_size, space_before, space_after):
    try:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True
    except Exception as e:
        print(f"Warning: Could not configure style {style_name}: {e}")

# --- Block rendering functions ---

def add_blockquote_block(doc, quote_lines):
    paragraphs_text = []
    current_para = []
    
    for line in quote_lines:
        cleaned = line.strip()
        if cleaned.startswith('>'):
            content = cleaned[1:].strip()
        else:
            content = cleaned
            
        if not content:
            if current_para:
                paragraphs_text.append(" ".join(current_para))
                current_para = []
        else:
            current_para.append(content)
            
    if current_para:
        paragraphs_text.append(" ".join(current_para))
        
    for text in paragraphs_text:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        parsed_inline = parse_inline_markdown(text)
        add_runs_to_paragraph(p, parsed_inline, default_font_size=10.5, default_font_name="Calibri", is_blockquote=True)

def add_code_block(doc, code_lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F5F7FA")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Border: 3pt left border in Navy
    tcPr = cell._tc.get_or_add_tcPr()
    existing_borders = tcPr.find(qn('w:tcBorders'))
    if existing_borders is not None:
        tcPr.remove(existing_borders)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="1B365D"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    
    code_text = "\n".join(code_lines)
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    cell.width = Inches(6.0)
    
    # Extra paragraph for clean spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def parse_markdown_table(lines):
    def split_row(line):
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        return [cell.strip() for cell in line.split('|')]
    
    headers = split_row(lines[0])
    separator = split_row(lines[1])
    
    alignments = []
    for sep in separator:
        if sep.startswith(':') and sep.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif sep.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            
    data_rows = []
    for line in lines[2:]:
        row_cells = split_row(line)
        if len(row_cells) < len(headers):
            row_cells += [''] * (len(headers) - len(row_cells))
        elif len(row_cells) > len(headers):
            row_cells = row_cells[:len(headers)]
        data_rows.append(row_cells)
        
    return headers, alignments, data_rows

def add_native_table(doc, headers, alignments, data_rows):
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Fill headers
    hdr_cells = table.rows[0].cells
    for col_idx, header_text in enumerate(headers):
        cell = hdr_cells[col_idx]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignments[col_idx] if col_idx < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        
        parsed_inline = parse_inline_markdown(header_text)
        add_runs_to_paragraph(paragraph, parsed_inline, default_font_size=10.5, default_font_name="Calibri")
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Fill data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F2F4F7" if row_idx % 2 == 1 else "FFFFFF"
        
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignments[col_idx] if col_idx < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            
            parsed_inline = parse_inline_markdown(cell_text)
            add_runs_to_paragraph(paragraph, parsed_inline, default_font_size=9.5, default_font_name="Calibri")
            
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
    table.autofit = True
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_list_item(doc, text, is_numbered=False):
    style_name = 'List Number' if is_numbered else 'List Bullet'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    parsed_inline = parse_inline_markdown(text)
    add_runs_to_paragraph(p, parsed_inline, default_font_size=11, default_font_name="Calibri")

# --- Compile single file to Word ---

def compile_single_markdown(input_path, output_path):
    doc = Document()
    
    # Configure Normal Style (Body text)
    body_style = doc.styles['Normal']
    body_style.font.name = 'Calibri'
    body_style.font.size = Pt(11)
    body_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.space_after = Pt(6)
    
    # Adjust List Styles
    for style_name in ['List Bullet', 'List Number']:
        try:
            lst_style = doc.styles[style_name]
            lst_style.font.name = 'Calibri'
            lst_style.font.size = Pt(11)
            lst_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            lst_style.paragraph_format.line_spacing = 1.15
            lst_style.paragraph_format.space_after = Pt(4)
        except Exception:
            pass
            
    # Configure Headings Styles
    configure_heading_style(doc, 'Heading 1', 18, 12, 6)
    configure_heading_style(doc, 'Heading 2', 14, 10, 4)
    configure_heading_style(doc, 'Heading 3', 12, 8, 2)
    
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    line_idx = 0
    total_lines = len(lines)
    
    while line_idx < total_lines:
        line = lines[line_idx]
        stripped = line.strip()
        
        if not stripped:
            line_idx += 1
            continue
            
        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("❖   ❖   ❖")
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            line_idx += 1
            continue
            
        # Headings
        h_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)
            level = min(level, 3)
            
            p_head = doc.add_heading('', level=level)
            p_head.text = "" # Clear default text
            
            parsed_inline = parse_inline_markdown(heading_text)
            font_size = 18 if level == 1 else (14 if level == 2 else 12)
            add_runs_to_heading(p_head, parsed_inline, font_size=font_size)
            
            line_idx += 1
            continue
            
        # Blockquotes
        if stripped.startswith('>'):
            quote_lines = []
            while line_idx < total_lines and (lines[line_idx].strip().startswith('>') or not lines[line_idx].strip()):
                if not lines[line_idx].strip():
                    peek_idx = line_idx + 1
                    is_still_bq = False
                    while peek_idx < total_lines:
                        p_strip = lines[peek_idx].strip()
                        if p_strip.startswith('>'):
                            is_still_bq = True
                            break
                        elif p_strip:
                            break
                        peek_idx += 1
                    
                    if is_still_bq:
                        quote_lines.append("")
                        line_idx = peek_idx
                        continue
                    else:
                        break
                
                quote_lines.append(lines[line_idx].strip())
                line_idx += 1
                
            add_blockquote_block(doc, quote_lines)
            continue
            
        # Fenced code / Equations block
        if stripped.startswith('```'):
            code_lines = []
            line_idx += 1
            while line_idx < total_lines and not lines[line_idx].strip().startswith('```'):
                code_lines.append(lines[line_idx].rstrip('\r\n'))
                line_idx += 1
                
            if line_idx < total_lines:
                line_idx += 1
                
            add_code_block(doc, code_lines)
            continue
            
        # Tables
        if stripped.startswith('|'):
            table_lines = []
            while line_idx < total_lines and lines[line_idx].strip().startswith('|'):
                table_lines.append(lines[line_idx].strip())
                line_idx += 1
                
            if len(table_lines) >= 2:
                headers, alignments, data_rows = parse_markdown_table(table_lines)
                add_native_table(doc, headers, alignments, data_rows)
            continue
            
        # Bullet list items
        if stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ '):
            bullet_content = stripped[2:]
            add_list_item(doc, bullet_content, is_numbered=False)
            line_idx += 1
            continue
            
        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            numbered_content = num_match.group(2)
            add_list_item(doc, numbered_content, is_numbered=True)
            line_idx += 1
            continue
            
        # Normal paragraphs
        p = doc.add_paragraph()
        parsed_inline = parse_inline_markdown(stripped)
        add_runs_to_paragraph(p, parsed_inline)
        line_idx += 1
        
    doc.save(output_path)
    print(f"  -> Generated: {os.path.basename(output_path)}")

# --- Main process logic ---

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Target folders
    word_en_dir = os.path.join(script_dir, "capitulos_word_en")
    word_pt_dir = os.path.join(script_dir, "capitulos_word_pt")
    
    os.makedirs(word_en_dir, exist_ok=True)
    os.makedirs(word_pt_dir, exist_ok=True)
    
    source_en_dir = os.path.join(script_dir, "capitulos")
    source_pt_dir = os.path.join(script_dir, "capitulos_pt")
    
    print("=" * 60)
    print("AVBC BILINGUAL SEPARATE CHAPTER WORD COMPILER")
    print("=" * 60)
    
    # English chapter processing
    if os.path.exists(source_en_dir):
        print(f"\nProcessing English chapters from '{source_en_dir}'...")
        en_files = [f for f in os.listdir(source_en_dir) if f.endswith('.md')]
        en_files.sort()
        for f in en_files:
            in_file = os.path.join(source_en_dir, f)
            out_file_name = f.replace('.md', '.docx')
            out_file = os.path.join(word_en_dir, out_file_name)
            compile_single_markdown(in_file, out_file)
    else:
        print(f"Error: English source folder not found at {source_en_dir}")
        
    # Portuguese chapter processing
    if os.path.exists(source_pt_dir):
        print(f"\nProcessing Portuguese chapters from '{source_pt_dir}'...")
        pt_files = [f for f in os.listdir(source_pt_dir) if f.endswith('.md')]
        pt_files.sort()
        for f in pt_files:
            in_file = os.path.join(source_pt_dir, f)
            out_file_name = f.replace('.md', '.docx')
            out_file = os.path.join(word_pt_dir, out_file_name)
            compile_single_markdown(in_file, out_file)
    else:
        print(f"Error: Portuguese source folder not found at {source_pt_dir}")
        
    print("\nCompilation completed successfully!")
    print(f"English separate chapters folder:   {word_en_dir}")
    print(f"Portuguese separate chapters folder: {word_pt_dir}")
    print("=" * 60)
