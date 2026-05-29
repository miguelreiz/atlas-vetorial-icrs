#!/usr/bin/env python3
"""
Compila todos os capítulos Markdown em um único Word Mestre (.docx)
com índice clicável, formatação profissional e navegação interna.

Uso:
  python3 compile_word.py pt   # Gera Word Mestre PT
  python3 compile_word.py en   # Gera Word Mestre EN
  python3 compile_word.py both # Gera ambos
"""

import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Configuração ───────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent

PT_CONFIG = {
    "title": "Biomecânica Vetorial",
    "subtitle": "Análise Vetorial Biomecânica Corneana\npara o Planejamento de Segmento de Anel Intraestromal",
    "author": "Dr. Miguel Reis\nDra. Jordana Sandes",
    "dir": BASE_DIR / "capitulos_pt",
    "output": BASE_DIR / "output" / "Biomecanica_Vetorial_PT.docx",
    "files": [
        "indice_geral.md",
        "cap01_cornea_biomecanica.md",
        "cap02_como_icrs_funcionam.md",
        "cap04_limites_nomogramas.md",       # Ch3 content
        "cap05_tres_dominios.md",            # Ch4 content
        "cap06_vetor_VR.md",                 # Ch5 content
        "cap07_vetor_VT.md",                 # Ch6 content
        "cap08_vetor_Vtau.md",               # Ch7 content
        "cap03_metodo_alpins.md",            # Ch8 content
        "cap09_classificacao_avbc.md",
        "cap11_fluxo_clinico.md",            # Ch10 content
        "cap12_casos_ilustrativos.md",       # Ch11 content
        "cap10_validacao_fem.md",            # Ch12 content
        "cap13_limitacoes_futuro.md",
        "cap14_plataforma_software.md",
        "cap15_conclusao.md",
        "apendice_A_modelagem_HGO.md",
        "apendice_B_scripts_febio.md",
        "apendice_C_glossario.md",
    ],
}

EN_CONFIG = {
    "title": "Corneal Biomechanical\nVector Analysis",
    "subtitle": "For Intracorneal Ring Segment Planning",
    "author": "Dr. Miguel Reis\nDr. Jordana Sandes",
    "dir": BASE_DIR / "capitulos_en",
    "output": BASE_DIR / "output" / "Corneal_Biomechanical_Vector_Analysis_EN.docx",
    "files": [
        "table_of_contents.md",
        "cap01_cornea_biomechanics.md",
        "cap02_how_icrs_work.md",
        "cap03_nomogram_limits.md",
        "cap04_three_domains.md",
        "cap05_vector_VR.md",
        "cap06_vector_VT.md",
        "cap07_vector_Vtau.md",
        "cap08_alpins_method.md",
        "cap09_avbc_classification.md",
        "cap11_clinical_workflow.md",
        "cap10_fem_validation.md",
        "cap12_case_studies.md",
        "cap13_limitations_future.md",
        "cap14_software_platform.md",
        "cap15_conclusion.md",
        "appendix_A_HGO_model.md",
        "appendix_B_febio_scripts.md",
        "appendix_C_glossary.md",
    ],
}


# ─── Estilos ────────────────────────────────────────────────────────
def setup_styles(doc):
    """Configura estilos profissionais do documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)  # Navy blue
    h1.paragraph_format.space_before = Pt(36)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x3A, 0x7C, 0xA5)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Quote style for blockquotes
    try:
        quote = doc.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        quote = doc.styles['BlockQuote']
    quote.font.name = 'Calibri'
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    quote.paragraph_format.left_indent = Cm(1.5)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)

    # Alert/Callout style
    try:
        alert = doc.styles.add_style('Alert', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        alert = doc.styles['Alert']
    alert.font.name = 'Calibri'
    alert.font.size = Pt(10.5)
    alert.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    alert.paragraph_format.left_indent = Cm(1.0)
    alert.paragraph_format.space_before = Pt(8)
    alert.paragraph_format.space_after = Pt(8)

    # Code style
    try:
        code = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code = doc.styles['CodeBlock']
    code.font.name = 'Consolas'
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    code.paragraph_format.left_indent = Cm(1.0)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)

    return doc


def add_cover_page(doc, config):
    """Adiciona capa profissional."""
    # Add several empty paragraphs for spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_lines = config["title"].split("\n")
    for line in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
        run.font.name = 'Calibri'

    doc.add_paragraph()

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

    doc.add_paragraph()

    # Subtitle
    sub_lines = config["subtitle"].split("\n")
    for line in sub_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(16)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["author"])
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = 'Calibri'

    # Year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Calibri'

    # Page break after cover
    doc.add_page_break()


def add_toc(doc):
    """Adiciona campo TOC (Table of Contents) do Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ÍNDICE")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Insert TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("(Clique com botão direito → Atualizar campo para gerar o índice)")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()


# ─── Parser Markdown → DOCX ────────────────────────────────────────
def parse_md_to_docx(doc, md_text, filename=""):
    """Converte texto Markdown para elementos DOCX."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    in_blockquote = False
    blockquote_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph(line, style='CodeBlock')
            i += 1
            continue

        # Skip YAML frontmatter
        if line.strip() == '---':
            i += 1
            continue

        # Empty line
        if not line.strip():
            # Flush blockquote
            if in_blockquote and blockquote_lines:
                flush_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
            # Flush table
            if in_table and table_rows:
                flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            stripped = line.strip()
            # Skip separator rows
            if re.match(r'^\|[\s\-:]+\|', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue

        # Flush table if we hit non-table line
        if in_table and table_rows:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Blockquotes
        if line.strip().startswith('>'):
            in_blockquote = True
            content = re.sub(r'^>\s*', '', line.strip())
            blockquote_lines.append(content)
            i += 1
            continue

        # Flush blockquote
        if in_blockquote and blockquote_lines:
            flush_blockquote(doc, blockquote_lines)
            blockquote_lines = []
            in_blockquote = False

        # Headings
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = clean_md(h_match.group(1 + 1))
            if level <= 3:
                style = f'Heading {level}'
                # First heading in file doesn't need page break if it's H1
                p = doc.add_heading(text, level=level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            text = clean_md(num_match.group(2))
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_md(line.strip()[2:])
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # LaTeX display math
        if line.strip().startswith('$$'):
            math_text = line.strip().replace('$$', '').strip()
            if not math_text:
                # Multi-line math
                i += 1
                math_lines = []
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    math_lines.append(lines[i])
                    i += 1
                math_text = ' '.join(math_lines)
                i += 1  # skip closing $$
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # Regular paragraph
        text = clean_md(line.strip())
        if text:
            p = doc.add_paragraph()
            add_formatted_text(p, text)

        i += 1

    # Flush any remaining
    if blockquote_lines:
        flush_blockquote(doc, blockquote_lines)
    if table_rows:
        flush_table(doc, table_rows)


def clean_md(text):
    """Remove Markdown formatting markers, keeping content."""
    # Remove inline math delimiters but keep content
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    # Remove image references
    text = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'[Figura: \1]', text)
    # Remove links, keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    """Adiciona texto com formatação inline (bold, italic)."""
    # Split on bold+italic, bold, italic patterns
    parts = re.split(r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def flush_blockquote(doc, lines):
    """Renderiza um bloco de citação/alerta."""
    text = ' '.join(lines)
    # Check for alert type
    alert_match = re.match(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*(.*)', text)
    if alert_match:
        alert_type = alert_match.group(1)
        content = alert_match.group(2).strip()
        # Remaining lines
        remaining = ' '.join(lines[1:]) if len(lines) > 1 else ''
        if remaining:
            content = content + ' ' + remaining if content else remaining

        # Alert header
        icons = {'NOTE': '📝', 'TIP': '💡', 'IMPORTANT': '⚠️', 'WARNING': '⚡', 'CAUTION': '🔴'}
        p = doc.add_paragraph(style='Alert')
        run = p.add_run(f"{icons.get(alert_type, '●')} {alert_type}")
        run.bold = True
        run.font.size = Pt(11)

        # Alert content
        if content:
            p2 = doc.add_paragraph(style='Alert')
            add_formatted_text(p2, clean_md(content))
    else:
        # Regular blockquote
        full_text = clean_md(text)
        if full_text.strip():
            p = doc.add_paragraph(full_text, style='BlockQuote')


def flush_table(doc, rows):
    """Renderiza uma tabela Markdown no DOCX."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = clean_md(cell_text.strip())
                # Header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)


# ─── Main ───────────────────────────────────────────────────────────
def compile_book(config):
    """Compila todos os capítulos em um único .docx."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_cover_page(doc, config)
    add_toc(doc)

    # Process each file
    chapter_dir = config["dir"]
    for filename in config["files"]:
        filepath = chapter_dir / filename
        if not filepath.exists():
            print(f"  ⚠ Ficheiro não encontrado: {filename}")
            continue

        print(f"  📄 Processando: {filename}")
        with open(filepath, 'r', encoding='utf-8') as f:
            md_text = f.read()
        parse_md_to_docx(doc, md_text, filename)

    # Save
    output_path = config["output"]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\n✅ Word Mestre salvo em: {output_path}")
    print(f"   Tamanho: {output_path.stat().st_size / 1024:.1f} KB")
    return output_path


if __name__ == "__main__":
    mode = sys.argv[1] if len(sys.argv) > 1 else "both"

    if mode in ("pt", "both"):
        print("\n🇧🇷 Compilando Word Mestre PT...")
        compile_book(PT_CONFIG)

    if mode in ("en", "both"):
        print("\n🇬🇧 Compilando Word Mestre EN...")
        compile_book(EN_CONFIG)

    print("\n🎉 Compilação concluída!")
    print("📋 Para gerar o índice clicável no Word:")
    print("   1. Abra o ficheiro .docx no Microsoft Word")
    print("   2. Clique no índice → Botão direito → 'Atualizar campo'")
    print("   3. Selecione 'Atualizar tabela inteira'")
    print("   4. Exporte como PDF: Ficheiro → Salvar como → PDF")
